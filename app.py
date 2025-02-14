import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
import nltk
from nltk.tokenize import sent_tokenize
from sklearn.feature_extraction.text import CountVectorizer
import matplotlib.pyplot as plt
from nltk.corpus import stopwords
import re
import base64
import io
from textblob import Word, TextBlob
from transformers import pipeline


nltk.download('stopwords')
stop_words = set(stopwords.words('english'))

nltk.download('punkt')


def extract_text_from_pdf(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text


def analyze_resume(text):
    words = re.findall(r'\b\w+\b', text.lower())
    filtered_words = [word for word in words if word not in stop_words]
    



    
    skills_list = [
        "python", "java", "c++", "html", "css", "javascript", "sql", "machine learning",
        "data science", "tensorflow", "excel", "operations", "team leadership", "management",
        "communication", "project planning", "problem solving", "design", "user experience",
        "front-end", "visual design", "user interface", "product management", "strategy",
        "business analysis", "leadership", "software development", "technical expertise",
        "data analysis", "data visualization", "statistics", "analytics", "research",
        "artificial intelligence", "big data", "cloud technologies", "data pipelines"
    ]
    skill_counts = {skill: filtered_words.count(skill) for skill in skills_list if skill in filtered_words}

    
    ranked_skills = {key: value for key, value in sorted(skill_counts.items(), key=lambda item: item[1], reverse=True)}

    
    experience = re.findall(r'(\d+)\s*(year|yrs|years)', text.lower())
    total_experience = sum(int(exp[0]) for exp in experience)

    return ranked_skills, total_experience


def extract_education_details(text):
    education_keywords = ["bachelor", "master", "degree", "university", "college", "phd", "engineering", "science"]
    education_details = []
    for line in text.splitlines():
        if any(keyword in line.lower() for keyword in education_keywords):
            education_details.append(line.strip())
    return education_details


def suggest_job_positions(skills):
    job_positions = {
        "Data Scientist": ["python", "machine learning", "data science", "tensorflow", "analytics", "statistics"],
        "Machine Learning Engineer": ["python", "machine learning", "tensorflow", "data science", "algorithm design"],
        "Software Developer": ["python", "java", "c++", "html", "css", "javascript", "software development", "mongodb", "php"],
        "Web Developer": ["html", "css", "javascript", "python", "web development", "frontend", "backend"],
        "AI Researcher": ["python", "machine learning", "data science", "tensorflow", "research", "artificial intelligence"],
        "Data Analyst": ["python", "data science", "sql", "machine learning", "data visualization", "statistics", "excel", "powerbi", "tableau"],
        "Project Manager": ["management", "team leadership", "communication", "project planning", "problem solving"],
        "Product Manager": ["product management", "leadership", "strategy", "communication", "project management"],
        "Business Analyst": ["business analysis", "data analysis", "communication", "problem solving", "project management"],
        "Technical Lead": ["leadership", "project management", "software development", "communication", "technical expertise"],
        "Data Engineer": ["python", "data science", "sql", "big data", "cloud technologies", "data pipelines"],
        "UX/UI Designer": ["design", "user experience", "front-end", "visual design", "user interface"],
        "Operations Manager": ["operations", "management", "team leadership", "communication", "problem solving", "research"],
        "Database Manager": ["sql", "database management", "data analysis", "communication", "problem solving", "mongodb"],
        "Entrepreneur": ["leadership", "strategy", "communication", "problem solving", "innovation", "finance", "business"]
    }

    suitable_jobs = []
    threshold = 0.5  
    for job, required_skills in job_positions.items():
        matching_skills = [skill for skill in required_skills if skill in skills]
        if len(matching_skills) / len(required_skills) >= threshold:
            suitable_jobs.append(job)

    return suitable_jobs


def display_resume(uploaded_file):
    if uploaded_file.type == "application/pdf":
        
        pdf_data = uploaded_file.read()
        st.download_button("Download Resume", data=pdf_data, file_name="resume.pdf", mime="application/pdf")
        st.markdown(f'<embed src="data:application/pdf;base64,{base64.b64encode(pdf_data).decode()}" width="100%" height="600px" type="application/pdf">', unsafe_allow_html=True)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        
        docx_data = uploaded_file.read()
        doc = Document(io.BytesIO(docx_data))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        st.text_area("Uploaded Resume", text, height=300)

@st.cache_resource
def load_nlp_model():
    """Load and cache the NLP model"""
    return pipeline(
        "text2text-generation",
        model="google/flan-t5-base",
        max_length=512,
        device="cpu"
    )

def perform_swot_analysis(text):
    """
    Perform SWOT analysis on resume text
    """
    
    sections = {
        'Strengths': [],
        'Weaknesses': [],
        'Opportunities': [],
        'Threats': []
    }

    try:
        sentences = sent_tokenize(text)
    except Exception as e:
        
        sentences = [s.strip() for s in text.split('.') if s.strip()]

    
    skill_indicators = [
        r'proficient in',
        r'expertise in',
        r'years? experience',
        r'successfully',
        r'achieved',
        r'led',
        r'managed',
        r'developed'
    ]
    
    weakness_indicators = [
        r'basic knowledge',
        r'familiar with',
        r'learning',
        r'limited experience',
        r'need to improve'
    ]

    opportunity_indicators = [
        r'goals',
        r'aim',
        r'aspire',
        r'seeking',
        r'interested in'
    ]

    
    for sentence in sentences:
        
        if any(re.search(pattern, sentence.lower()) for pattern in skill_indicators):
            sections['Strengths'].append(sentence.strip())

        
        if any(keyword in sentence.lower() for keyword in ['goals', 'aim', 'aspire', 'seeking', 'interested in']):
            sections['Opportunities'].append(sentence.strip())

        
        if any(keyword in sentence.lower() for keyword in ['basic knowledge', 'familiar with', 'learning']):
            sections['Weaknesses'].append(sentence.strip())

    
    sections['Threats'] = [
        "Rapidly evolving technology landscape requiring continuous learning",
        "Competitive job market with increasing skill requirements",
        "Need to stay updated with industry trends"
    ]

    return sections

def display_swot_analysis(sections):
    """
    Display SWOT analysis results
    """
    if not sections:
        st.error("Unable to generate SWOT analysis")
        return

    st.subheader("SWOT Analysis")

    col1, col2 = st.columns(2)

    with col1:
        
        st.markdown("### 💪 Strengths")
        with st.container():
            st.markdown("""<div style='background-color: #e6ffe6; padding: 15px; border-radius: 5px;'>""",
                      unsafe_allow_html=True)
            for strength in sections['Strengths']:
                st.markdown(f"• {strength}")
            st.markdown("</div>", unsafe_allow_html=True)

        
        st.markdown("### 🎯 Opportunities")
        with st.container():
            st.markdown("""<div style='background-color: #e6f3ff; padding: 15px; border-radius: 5px;'>""",
                      unsafe_allow_html=True)
            for opportunity in sections['Opportunities']:
                st.markdown(f"• {opportunity}")
            st.markdown("</div>", unsafe_allow_html=True)

    with col2:
        
        st.markdown("### ⚠️ Areas for Improvement")
        with st.container():
            st.markdown("""<div style='background-color: #fff3e6; padding: 15px; border-radius: 5px;'>""",
                      unsafe_allow_html=True)
            for weakness in sections['Weaknesses']:
                st.markdown(f"• {weakness}")
            st.markdown("</div>", unsafe_allow_html=True)

        
        st.markdown("### ⚡ Challenges")
        with st.container():
            st.markdown("""<div style='background-color: #ffe6e6; padding: 15px; border-radius: 5px;'>""",
                      unsafe_allow_html=True)
            for threat in sections['Threats']:
                st.markdown(f"• {threat}")
            st.markdown("</div>", unsafe_allow_html=True)


st.title("ResuMind- An AI-Based Resume Analyzer")
st.write("Upload a resume (PDF or DOCX) to analyze the candidate's skills and experience.")

uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx"])

if uploaded_file:
    
    display_resume(uploaded_file)

    
    if uploaded_file.type == "application/pdf":
        resume_text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        resume_text = extract_text_from_docx(uploaded_file)

    
    skills, experience = analyze_resume(resume_text)

    
    st.subheader("Skills")
    skill_names = list(skills.keys())
    skill_counts = list(skills.values())


    st.subheader("Total Experience")
    st.write(f"{experience} years of experience")

    
    st.subheader("Suggested Job Positions Based on Skills")
    job_positions = suggest_job_positions(skills)
    if job_positions:
        for job in job_positions:
            st.write(f"- {job}")
    else:
        st.write("No specific job suggestions found based on the resume's skills.")

    
    st.subheader("Educational Details")
    education_details = extract_education_details(resume_text)
    if education_details:
        for edu in education_details:
            st.write(f"- {edu}")
    else:
        st.write("No educational details found.")

    
    st.subheader("Personal Information")
    personal_info = re.findall(r'(Name|Email|Phone):\s*(.*)', resume_text)
    if personal_info:
        for info in personal_info:
            st.write(f"{info[0]}: {info[1]}")
    else:
        st.write("No personal information found.")

    
    st.subheader("Keyword Highlighting")
    keywords = [
        "python", "java", "c++", "html", "css", "javascript", "sql", "machine learning",
        "data science", "tensorflow", "excel", "operations", "team leadership", "management",
        "communication", "project planning", "problem solving", "design", "user experience",
        "front-end", "visual design", "user interface", "product management", "strategy",
        "business analysis", "leadership", "software development", "technical expertise",
        "data analysis", "data visualization", "statistics", "analytics", "research",
        "artificial intelligence", "big data", "cloud technologies", "data pipelines"
    ]
    highlighted_text = resume_text
    for keyword in keywords:
        highlighted_text = re.sub(r'\b' + re.escape(keyword) + r'\b', f"{keyword}", highlighted_text, flags=re.IGNORECASE)
    st.markdown(highlighted_text)

    st.write("Generating SWOT Analysis...")
    with st.spinner("Analyzing resume content..."):
        swot_sections = perform_swot_analysis(resume_text)
        if swot_sections:
            display_swot_analysis(swot_sections)

    
    st.subheader("Sentiment Analysis")
    analysis = TextBlob(resume_text)
    sentiment = analysis.sentiment
    st.write(f"Polarity: {sentiment.polarity}, Subjectivity: {sentiment.subjectivity}")

    
    st.subheader("Grammar and Spell Check")
    words = resume_text.split()
    corrected_words = [str(Word(word).correct()) for word in words]
    corrected_text = ' '.join(corrected_words)
    st.text_area("Corrected Resume", corrected_text, height=300)

    
    st.subheader("Personality Traits Analysis")
    personality_traits = {
        "Leadership": ["leadership", "management", "team"],
        "Creativity": ["creative", "innovative", "design"],
        "Analytical": ["analytical", "data", "analysis"],
        "Communication": ["communication", "presentation", "writing"]
    }
    traits_found = {trait: any(keyword in resume_text.lower() for keyword in keywords) for trait, keywords in personality_traits.items()}
    for trait, found in traits_found.items():
        st.write(f"{trait}: {'Found' if found else 'Not Found'}")

    
    st.subheader("Visualization Dashboard")
    st.write("This section provides a visual overview of the resume analysis.")

    
    st.subheader("Skill Distribution")
    plt.figure(figsize=(8, 8))
    plt.pie(skill_counts, labels=skill_names, autopct='%1.1f%%', startangle=140)
    plt.title("Skill Distribution")
    st.pyplot(plt)

    
    st.subheader("Experience Timeline")
    experience_timeline = re.findall(r'(\d+)\s*(year|yrs|years)\s*at\s*(.*)', resume_text.lower())
    if experience_timeline:
        for exp in experience_timeline:
            st.write(f"{exp[0]} years at {exp[2]}")
    else:
        st.write("No experience timeline found.")