import streamlit as st
from PyPDF2 import PdfReader
from docx import Document
import nltk
from sklearn.feature_extraction.text import CountVectorizer
import matplotlib.pyplot as plt
from nltk.corpus import stopwords
import re
import base64
import io
from textblob import Word, TextBlob
from transformers import pipeline


# Download NLTK stopwords (for removing common words)
nltk.download('stopwords')
stop_words = set(stopwords.words('english'))

# Initialize the LLM pipeline
# llm_pipeline = pipeline("text-generation", model="gpt-2")

@st.cache_resource
def load_llm():
    return pipeline(
        "text2text-generation",
        model="google/flan-t5-base",  # Using FLAN-T5 for better reasoning capabilities
        max_length=512,
        device="cpu"
    )


# Function to extract text from a PDF
def extract_text_from_pdf(pdf_file):
    pdf_reader = PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# Function to extract text from Word document
def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    return text

# Function to process the text and analyze skills
def analyze_resume(text):
    words = re.findall(r'\b\w+\b', text.lower())
    filtered_words = [word for word in words if word not in stop_words]

    # Example skills list to look for
    skills_list = [
        "python", "java", "c++", "html", "css", "javascript", "sql", "machine learning",
        "data science", "tensorflow", "excel", "operations", "team leadership", "management",
        "communication", "project planning", "problem solving", "design", "user experience",
        "front-end", "visual design", "user interface", "product management", "strategy",
        "business analysis", "leadership", "software development", "technical expertise",
        "data analysis", "data visualization", "statistics", "analytics", "research",
        "artificial intelligence", "big data", "cloud technologies", "data pipelines"
    ]
    skill_counts = {skill: filtered_words.count(skill) for skill in skills_list if skill in filtered_words}

    # Rank skills by count
    ranked_skills = {key: value for key, value in sorted(skill_counts.items(), key=lambda item: item[1], reverse=True)}

    # Extract years of experience (very basic, assuming a format like "5 years")
    experience = re.findall(r'(\d+)\s*(year|yrs|years)', text.lower())
    total_experience = sum(int(exp[0]) for exp in experience)

    return ranked_skills, total_experience

# Function to extract education details (very basic pattern recognition)
def extract_education_details(text):
    education_keywords = ["bachelor", "master", "degree", "university", "college", "phd", "engineering", "science"]
    education_details = []
    for line in text.splitlines():
        if any(keyword in line.lower() for keyword in education_keywords):
            education_details.append(line.strip())
    return education_details

# Function to determine suitable job positions based on skills
def suggest_job_positions(skills):
    job_positions = {
        "Data Scientist": ["python", "machine learning", "data science", "tensorflow", "analytics", "statistics"],
        "Machine Learning Engineer": ["python", "machine learning", "tensorflow", "data science", "algorithm design"],
        "Software Developer": ["python", "java", "c++", "html", "css", "javascript", "software development", "mongodb", "php"],
        "Web Developer": ["html", "css", "javascript", "python", "web development", "frontend", "backend"],
        "AI Researcher": ["python", "machine learning", "data science", "tensorflow", "research", "artificial intelligence"],
        "Data Analyst": ["python", "data science", "sql", "machine learning", "data visualization", "statistics", "excel", "powerbi", "tableau"],
        "Project Manager": ["management", "team leadership", "communication", "project planning", "problem solving"],
        "Product Manager": ["product management", "leadership", "strategy", "communication", "project management"],
        "Business Analyst": ["business analysis", "data analysis", "communication", "problem solving", "project management"],
        "Technical Lead": ["leadership", "project management", "software development", "communication", "technical expertise"],
        "Data Engineer": ["python", "data science", "sql", "big data", "cloud technologies", "data pipelines"],
        "UX/UI Designer": ["design", "user experience", "front-end", "visual design", "user interface"],
        "Operations Manager": ["operations", "management", "team leadership", "communication", "problem solving", "research"],
        "Database Manager": ["sql", "database management", "data analysis", "communication", "problem solving", "mongodb"],
        "Entrepreneur": ["leadership", "strategy", "communication", "problem solving", "innovation", "finance", "business"]
    }

    suitable_jobs = []
    threshold = 0.5  # Adjust the threshold as needed
    for job, required_skills in job_positions.items():
        matching_skills = [skill for skill in required_skills if skill in skills]
        if len(matching_skills) / len(required_skills) >= threshold:
            suitable_jobs.append(job)

    return suitable_jobs

# Function to display the resume (for PDF or DOCX files)
def display_resume(uploaded_file):
    if uploaded_file.type == "application/pdf":
        # Display PDF
        pdf_data = uploaded_file.read()
        st.download_button("Download Resume", data=pdf_data, file_name="resume.pdf", mime="application/pdf")
        st.markdown(f'<embed src="data:application/pdf;base64,{base64.b64encode(pdf_data).decode()}" width="100%" height="600px" type="application/pdf">', unsafe_allow_html=True)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        # Display DOCX content
        docx_data = uploaded_file.read()
        doc = Document(io.BytesIO(docx_data))
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        st.text_area("Uploaded Resume", text, height=300)

# Function to extract personal information using regular expressions
def extract_personal_info(text):
    personal_info = {}

    # Extract name (supports multiple formats)
    name_patterns = [
        r'(?i)(?:name|full name)\s*:?\s*([\w\s]+)',
        r'^([\w\s]+)$',  # Name at the start of resume
        r'(?i)(?:i am|this is)\s+([\w\s]+)'
    ]
    
    for pattern in name_patterns:
        name_match = re.search(pattern, text)
        if name_match:
            name = name_match.group(1).strip()
            if len(name.split()) >= 2:  # Ensure it's a full name
                personal_info['Name'] = name
                break

    # Extract email (supports multiple email formats)
    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
    email_matches = re.findall(email_pattern, text)
    if email_matches:
        personal_info['Email'] = email_matches[0].strip()

    # Extract phone numbers (supports multiple formats)
    phone_patterns = [
        r'(?:(?:\+\d{1,3}[-.\s]?)?(?:\d{3}[-.\s]?\d{3}[-.\s]?\d{4}))',  # +1 123-456-7890
        r'(?:\d{3}[-.\s]?\d{3}[-.\s]?\d{4})',  # 123-456-7890
        r'(?:\+\d{1,3}\s)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # (123) 456-7890
        r'\b\d{10}\b',  # 1234567890
        r'(?:\+\d{1,3}[-.\s]?)?\d{3,}[-.\s]?\d{3,}[-.\s]?\d{3,}'  # International formats
    ]
    
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            # Clean up the phone number format
            phone = re.sub(r'[-.\s\(\)]', '', phone_match.group(0))
            if len(phone) >= 10:  # Ensure it's a valid length
                personal_info['Phone'] = phone
                break

    # Extract location/address
    location_patterns = [
        r'(?i)(?:address|location|city|residing at)\s*:?\s*([\w\s,.-]+)',
        r'(?i)(?:residing|based) in\s+([\w\s,.-]+)',
        r'(?i)([\w\s,.-]+(?:street|avenue|road|boulevard|lane|drive|city|state|zip|pincode)[\w\s,.-]+)'
    ]
    
    for pattern in location_patterns:
        location_match = re.search(pattern, text)
        if location_match:
            personal_info['Location'] = location_match.group(1).strip()
            break

    # Extract LinkedIn profile
    linkedin_pattern = r'(?i)(?:linkedin\.com/in/|linkedin:?\s*)?(?:https?://)?(?:www\.)?linkedin\.com/in/([\w-]+)/?'
    linkedin_match = re.search(linkedin_pattern, text)
    if linkedin_match:
        personal_info['LinkedIn'] = f"linkedin.com/in/{linkedin_match.group(1)}"

    # Extract website/portfolio
    website_pattern = r'(?i)(?:website|portfolio|blog)\s*:?\s*((?:https?://)?[\w.-]+(?:\.[\w.-]+)+[\w\-._~:/?#[\]@!$&\'()*+,;=]*)'
    website_match = re.search(website_pattern, text)
    if website_match:
        personal_info['Website'] = website_match.group(1)

    # Extract additional professional details
    professional_patterns = {
        'Title': r'(?i)(?:designation|position|title|role)\s*:?\s*([\w\s]+)',
        'Industry': r'(?i)(?:industry|sector)\s*:?\s*([\w\s]+)',
        'Experience': r'(?i)(\d+)\+?\s*(?:years? of experience|years? in|yrs?)',
    }
    
    for key, pattern in professional_patterns.items():
        match = re.search(pattern, text)
        if match:
            personal_info[key] = match.group(1).strip()

    return personal_info

# Function to display personal information in a structured format
def display_personal_info(personal_info):
    if not personal_info:
        st.write("No personal information found.")
        return

    # Create two columns for better layout
    col1, col2 = st.columns(2)

    # Essential info in first column
    with col1:
        if 'Name' in personal_info:
            st.markdown(f"**👤 Name:** {personal_info['Name']}")
        if 'Email' in personal_info:
            st.markdown(f"**📧 Email:** {personal_info['Email']}")
        if 'Phone' in personal_info:
            st.markdown(f"**📱 Phone:** {personal_info['Phone']}")
        if 'Location' in personal_info:
            st.markdown(f"**📍 Location:** {personal_info['Location']}")

    # Professional info in second column
    with col2:
        if 'Title' in personal_info:
            st.markdown(f"**💼 Title:** {personal_info['Title']}")
        if 'Experience' in personal_info:
            st.markdown(f"**⏳ Experience:** {personal_info['Experience']} years")
        if 'LinkedIn' in personal_info:
            st.markdown(f"**🔗 LinkedIn:** [{personal_info['LinkedIn']}](https://{personal_info['LinkedIn']})")
        if 'Website' in personal_info:
            st.markdown(f"**🌐 Website:** [{personal_info['Website']}]({personal_info['Website']})")




# Function to perform SWOT analysis using LLM
def perform_swot_analysis(resume_text, llm):
    # Create a structured prompt for better analysis
    prompt = f"""
    Analyze this resume and provide a detailed SWOT analysis:
    Resume: {resume_text}
    
    Task: Generate a structured SWOT (Strengths, Weaknesses, Opportunities, Threats) analysis 
    focusing on the candidate's professional profile.
    
    Format your response as follows:
    Strengths:
    - List key technical skills and experiences
    - Highlight demonstrated achievements
    - Note relevant qualifications
    
    Weaknesses:
    - Identify gaps in skills or experience
    - Point out missing relevant certifications
    - Note areas needing improvement
    
    Opportunities:
    - Suggest potential career growth paths
    - Identify emerging industry trends matching skills
    - Recommend skill development areas
    
    Threats:
    - Consider industry changes affecting role
    - Note competitive job market factors
    - Identify potential skill obsolescence
    """
    
    try:
        # Generate the analysis
        response = llm(prompt, max_length=1024, num_return_sequences=1)[0]['generated_text']
        
        # Process and structure the response
        sections = {
            'Strengths': [],
            'Weaknesses': [],
            'Opportunities': [],
            'Threats': []
        }
        
        current_section = None
        for line in response.split('\n'):
            line = line.strip()
            if line.lower().startswith(('strengths:', 'weaknesses:', 'opportunities:', 'threats:')):
                current_section = line.split(':')[0].strip()
            elif line.startswith('-') and current_section:
                sections[current_section].append(line.strip('- '))
        
        return sections
        
    except Exception as e:
        st.error(f"Error generating SWOT analysis: {str(e)}")
        return None
def display_swot_analysis(sections):
        if not sections:
          st.error("Unable to generate SWOT analysis")
        return
        
        st.subheader("AI-Generated SWOT Analysis")
    
    # Create two columns for the SWOT analysis
        col1, col2 = st.columns(2)
    
        with col1:
        # Strengths (Green)
        st.markdown("### 💪 Strengths")
        st.markdown("""
        <div style='background-color: #e6ffe6; padding: 15px; border-radius: 5px;'>
        """, unsafe_allow_html=True)
        for strength in sections['Strengths']:
            st.markdown(f"• {strength}")
        st.markdown("</div>", unsafe_allow_html=True)
        
        # Opportunities (Blue)
        st.markdown("### 🎯 Opportunities")
        st.markdown("""
        <div style='background-color: #e6f3ff; padding: 15px; border-radius: 5px;'>
        """, unsafe_allow_html=True)
        for opportunity in sections['Opportunities']:
            st.markdown(f"• {opportunity}")
        st.markdown("</div>", unsafe_allow_html=True)
    
    with col2:
        # Weaknesses (Yellow)
        st.markdown("### ⚠️ Weaknesses")
        st.markdown("""
        <div style='background-color: #fff3e6; padding: 15px; border-radius: 5px;'>
        """, unsafe_allow_html=True)
        for weakness in sections['Weaknesses']:
            st.markdown(f"• {weakness}")
        st.markdown("</div>", unsafe_allow_html=True)
        
        # Threats (Red)
        st.markdown("### ⚡ Threats")
        st.markdown("""
        <div style='background-color: #ffe6e6; padding: 15px; border-radius: 5px;'>
        """, unsafe_allow_html=True)
        for threat in sections['Threats']:
            st.markdown(f"• {threat}")
        st.markdown("</div>", unsafe_allow_html=True)


#  UI
st.title("AI-Based Resume Analyzer for Job Seekers")
st.write("Upload a resume (PDF or DOCX) to analyze the candidate's skills and experience.")

uploaded_file = st.file_uploader("Choose a file", type=["pdf", "docx"])

if uploaded_file:
    # Display the uploaded resume
    display_resume(uploaded_file)

    # Extract text based on file type
    if uploaded_file.type == "application/pdf":
        resume_text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        resume_text = extract_text_from_docx(uploaded_file)

    # Analyze the resume
    skills, experience = analyze_resume(resume_text)
    education_details = extract_education_details(resume_text)

    # Display skills with visualization (Bar Chart)
    st.subheader("Skills")
    skill_names = list(skills.keys())
    skill_counts = list(skills.values())


    # Display total experience
    st.subheader("Total Experience")
    st.write(f"{experience} years of experience")

    # Display job suggestions
    st.subheader("Suggested Job Positions Based on Skills")
    job_positions = suggest_job_positions(skills)
    if job_positions:
        for job in job_positions:
            st.write(f"- {job}")
    else:
        st.write("No specific job suggestions found based on the resume's skills.")

    # Extract and display educational details
    st.subheader("Educational Details")
    if education_details:
        for edu in education_details:
            st.write(f"- {edu}")
    else:
        st.write("No educational details found.")

    # Display personal information
    st.subheader("Personal Information")
    personal_info = extract_personal_info(resume_text)
    display_personal_info(personal_info)

    # Display keyword highlighting
    st.subheader("Keyword Highlighting")
    keywords = [
        "python", "java", "c++", "html", "css", "javascript", "sql", "machine learning",
        "data science", "tensorflow", "excel", "operations", "team leadership", "management",
        "communication", "project planning", "problem solving", "design", "user experience",
        "front-end", "visual design", "user interface", "product management", "strategy",
        "business analysis", "leadership", "software development", "technical expertise",
        "data analysis", "data visualization", "statistics", "analytics", "research",
        "artificial intelligence", "big data", "cloud technologies", "data pipelines"
    ]
    highlighted_text = resume_text
    for keyword in keywords:
        highlighted_text = re.sub(r'\b' + re.escape(keyword) + r'\b', f"**{keyword}**", highlighted_text, flags=re.IGNORECASE)
    st.markdown(highlighted_text)

    # Display sentiment analysis
    st.subheader("Sentiment Analysis")
    analysis = TextBlob(resume_text)
    sentiment = analysis.sentiment
    st.write(f"Polarity: {sentiment.polarity}, Subjectivity: {sentiment.subjectivity}")

    # Display grammar and spell check
    st.subheader("Grammar and Spell Check")
    words = resume_text.split()
    corrected_words = [str(Word(word)).correct() for word in words]
    corrected_text = ' '.join(corrected_words)
    st.text_area("Corrected Resume", corrected_text, height=300)

    # Display personality traits analysis (basic example)
    st.subheader("Personality Traits Analysis")
    personality_traits = {
        "Leadership": ["leadership", "management", "team"],
        "Creativity": ["creative", "innovative", "design"],
        "Analytical": ["analytical", "data", "analysis"],
        "Communication": ["communication", "presentation", "writing"]
    }
    traits_found = {trait: any(keyword in resume_text.lower() for keyword in keywords) for trait, keywords in personality_traits.items()}
    for trait, found in traits_found.items():
        st.write(f"{trait}: {'Found' if found else 'Not Found'}")

    # Visualization Dashboard
    st.subheader("Visualization Dashboard")
    st.write("This section provides a visual overview of the resume analysis.")

    # Skill Distribution Pie Chart
    st.subheader("Skill Distribution")
    plt.figure(figsize=(8, 8))
    plt.pie(skill_counts, labels=skill_names, autopct='%1.1f%%', startangle=140)
    plt.title("Skill Distribution")
    st.pyplot(plt)

    # Experience Timeline (basic example)
    st.subheader("Experience Timeline")
    experience_timeline = re.findall(r'(\d+)\s*(year|yrs|years)\s*at\s*(.*)', resume_text.lower())
    if experience_timeline:
        for exp in experience_timeline:
            st.write(f"{exp[0]} years at {exp[2]}")
    else:
        st.write("No experience timeline found.")

if uploaded_file:

    # Load the LLM model
    llm = load_llm()
    
    # Perform SWOT analysis
    st.write("Generating AI-powered SWOT analysis...")
    with st.spinner("This may take a few moments..."):
        swot_sections = perform_swot_analysis(resume_text, llm)
        display_swot_analysis(swot_sections)
